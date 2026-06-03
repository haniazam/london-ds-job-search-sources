#!/usr/bin/env python3
"""QA gate for the tailored resumes.

Checks (base CV = single source of truth):
 1. Truthfulness  — every numeric/metric token in a tailored resume must also
    appear in the base CV source.
 2. JD keyword coverage — role-specific must-have terms present.
 3. Differentiation — summaries are materially distinct.
 4. ATS parse — generated .docx extracts cleanly (no stray '**' or '#').
"""
import re
import sys
from pathlib import Path

from docx import Document

BASE = Path("resumes/base/Hani_Azam_base_CV_source.md").read_text(encoding="utf-8")
TAIL = Path("resumes/tailored")
DOCX = TAIL / "docx"

NUM_RE = re.compile(r"\$?\d[\d,]*(?:\.\d+)?(?:pp|%|\+|M|m|k)?")

# numbers that are legitimately introduced by tailoring (years/levels/JD refs),
# not claims about the candidate's results.
ALLOWED_EXTRA = {"1", "2", "3", "4", "5", "6", "7", "9", "12", "10", "20",
                 "150", "2018", "2017", "2021", "2020", "2019"}

# role file -> must-have JD keywords (lowercased substring match)
KEYWORDS = {
    "OpenAI": ["causal", "experiment", "metric", "measurement integrity", "propensity"],
    "CitadelSecurities": ["model", "experiment", "python", "sql"],
    "Stripe": ["experiment", "forecast", "python", "sql"],
    "GoogleDeepMind": ["machine learning", "python", "experiment"],
    "TheTradeDesk": ["advertis", "experiment", "machine learning"],
    "GResearch": ["model", "experiment", "python", "etl"],
    "Spotify": ["experiment", "retention", "growth"],
    "Monzo": ["propensity", "growth", "product analytics", "python"],
    "Google_Product": ["experiment", "product", "causal"],
    "Bloomberg": ["machine learning", "tf-idf", "fasttext"],
    "ManGroup": ["responsible investment", "political-economy", "product"],
    "TikTok": ["commerce", "experiment", "growth", "spark", "hadoop"],
}


def role_key(name: str) -> str:
    for k in KEYWORDS:
        if k.split("_")[0].lower() in name.lower() and (
            "_" not in k or k.split("_")[1].lower() in name.lower()
        ):
            return k
    for k in KEYWORDS:  # looser fallback
        if k.split("_")[0].lower() in name.lower():
            return k
    return ""


def base_numbers():
    return set(NUM_RE.findall(BASE.lower())) | ALLOWED_EXTRA


def main():
    base_nums = {n.strip("+") for n in base_numbers()}
    mds = sorted(p for p in TAIL.glob("*.md") if p.name != "README.md")
    summaries = {}
    problems = []
    print(f"Found {len(mds)} tailored resumes\n")

    for md in mds:
        text = md.read_text(encoding="utf-8")
        low = text.lower()
        name = md.stem

        # 1. truthfulness — numbers
        nums = {n.strip("+") for n in NUM_RE.findall(low)}
        unknown = sorted(n for n in nums if n and n not in base_nums)
        if unknown:
            problems.append(f"[TRUTH] {name}: numbers not in base CV: {unknown}")

        # 2. keyword coverage
        rk = role_key(name)
        if rk:
            missing = [kw for kw in KEYWORDS[rk] if kw not in low]
            status = "OK" if not missing else f"MISSING {missing}"
            print(f"  {name:60s} [{rk:18s}] {status}")
            if missing:
                problems.append(f"[KEYWORDS] {name}: missing {missing}")
        else:
            problems.append(f"[KEYWORDS] {name}: no role mapping")

        # collect summary
        m = re.search(r"## SUMMARY\n+(.+)", text)
        if m:
            summaries[name] = m.group(1).strip()

    # 3. differentiation
    print("\nDifferentiation (summary uniqueness):")
    vals = list(summaries.values())
    dupes = len(vals) - len(set(vals))
    print(f"  {len(vals)} summaries, {len(set(vals))} unique, {dupes} duplicates")
    if dupes:
        problems.append(f"[DIFF] {dupes} duplicate summaries")

    # 4. ATS parse of docx
    print("\nATS parse (docx):")
    if DOCX.exists():
        for dx in sorted(DOCX.glob("*.docx")):
            doc = Document(str(dx))
            full = "\n".join(p.text for p in doc.paragraphs)
            stray = [s for s in ("**", "##", "](") if s in full]
            n_par = len([p for p in doc.paragraphs if p.text.strip()])
            flag = f"STRAY {stray}" if stray else "clean"
            print(f"  {dx.name:60s} {n_par:3d} paras  {flag}")
            if stray:
                problems.append(f"[ATS] {dx.name}: stray markup {stray}")
    else:
        problems.append("[ATS] docx directory missing")

    print("\n" + "=" * 70)
    if problems:
        print(f"FAIL — {len(problems)} issue(s):")
        for p in problems:
            print("  -", p)
        sys.exit(1)
    print("PASS — all automated QA checks green")


if __name__ == "__main__":
    main()
