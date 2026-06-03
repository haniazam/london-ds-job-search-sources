#!/usr/bin/env python3
"""Convert the tailored resume markdown files into formatted .docx documents.

Handles the limited markdown subset used in resumes/tailored/*.md:
  # H1            -> centered name
  ## H2           -> section heading (bold, spaced)
  - bullet        -> bullet list item (with inline **bold** / *italic*)
  **Tailored...** -> normal paragraph
  plain text      -> normal paragraph
Inline: **bold**, *italic*, [text](url) -> text.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")
TOKEN_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def add_runs(paragraph, text):
    """Add inline-formatted runs (handles **bold** and *italic*)."""
    text = LINK_RE.sub(r"\1", text)
    for part in TOKEN_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def convert(md_path: Path, out_path: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(18)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.space_before = Pt(10)
            run = p.add_run(line[3:].strip().upper())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x5C)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            # Center the contact line (second line, contains the email)
            if "azamhani95@gmail.com" in line:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, line.strip())

    doc.save(str(out_path))
    print(f"wrote {out_path}")


def main():
    src = Path("resumes/tailored")
    out = Path("resumes/tailored/docx")
    out.mkdir(parents=True, exist_ok=True)
    for md in sorted(src.glob("*.md")):
        if md.name == "README.md":
            continue
        convert(md, out / (md.stem + ".docx"))


if __name__ == "__main__":
    main()
